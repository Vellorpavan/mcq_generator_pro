from docx import Document
from docx.shared import Pt
import io
import re
from typing import List
from app.models.mcq import MCQItem

def clean_option(text: str) -> str:
    return re.sub(r'^[\(\[]?\s*[A-E]\s*[\.\)\]]\s*', '', text, flags=re.IGNORECASE).strip()

def write_cell(cell, text: str, bold: bool = False):
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(11)

def generate_docx(mcq_list: List[MCQItem], doc_title: str, mode: str = "full") -> io.BytesIO:
    doc = Document()
    
    if mode == "key":
        doc.add_heading(doc_title.upper() if doc_title else "ANSWER KEY", level=1)
        for i, mcq in enumerate(mcq_list, start=1):
            ans = mcq.correct_answer or "N/A"
            doc.add_paragraph(f"{i}. {ans}")
    else:
        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"

        # Header
        hdr = table.rows[0].cells
        hdr[0].text = "S.No"
        hdr[1].text = doc_title.upper() if doc_title else "CHOOSE THE CORRECT ANSWERS:"
        hdr[1].merge(hdr[2])
        hdr[1].merge(hdr[3])

        # Questions
        for i, mcq in enumerate(mcq_list, start=1):
            # Question row
            qrow = table.add_row().cells
            qrow[0].text = f"{i}."
            qrow[1].text = mcq.question
            qrow[1].merge(qrow[2])
            qrow[3].text = ""

            opts = [clean_option(o) for o in mcq.options]
            ans = mcq.correct_answer

            # mode == no_answers never bolds
            show_ans = (mode == "full")

            # A | B
            r1 = table.add_row().cells
            r1[0].merge(qrow[0])
            write_cell(r1[1], f"A) {opts[0]}", bold=(show_ans and ans == "A"))
            write_cell(r1[2], f"B) {opts[1]}", bold=(show_ans and ans == "B"))
            r1[3].merge(qrow[3])

            # C | D
            r2 = table.add_row().cells
            r2[0].merge(qrow[0])
            write_cell(r2[1], f"C) {opts[2]}", bold=(show_ans and ans == "C"))
            write_cell(r2[2], f"D) {opts[3]}", bold=(show_ans and ans == "D"))
            r2[3].merge(qrow[3])

    # Save to memory
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream
